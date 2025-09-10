import json
import os
import re
from docx import Document
from openpyxl import Workbook

FILE_NAME = "danhba.json"
PAGE_SIZE = 5


# ================== HÀM XỬ LÝ ==================

def load_contacts():
    if not os.path.exists(FILE_NAME):
        return []
    with open(FILE_NAME, "r", encoding="utf-8") as f:
        return json.load(f)


def save_contacts(contacts):
    with open(FILE_NAME, "w", encoding="utf-8") as f:
        json.dump(contacts, f, ensure_ascii=False, indent=4)


def normalize_phone(phone: str) -> str:
    phone = re.sub(r"\D", "", phone)
    return phone


def validate_phone(phone: str) -> bool:
    return bool(re.fullmatch(r"\d{10}", phone))


def validate_email(email: str) -> bool:
    if not email:
        return True
    return bool(re.fullmatch(r"[^@]+@[^@]+\.[^@]+", email))


def input_optional(prompt: str, default="") -> str:
    value = input(prompt).strip()
    return value if value else default


# ================== QUẢN LÝ DANH BẠ ==================

def add_contact(contacts):
    name = input("Nhập tên: ").strip()
    phone = normalize_phone(input("Nhập số điện thoại (10 số): ").strip())
    if not validate_phone(phone):
        print("❌ Số điện thoại không hợp lệ!\n")
        return
    if any(c['phone'] == phone for c in contacts):
        print("❌ Số điện thoại đã tồn tại!\n")
        return

    email = input_optional("Nhập email (có thể bỏ trống): ")
    if not validate_email(email):
        print("❌ Email không hợp lệ!\n")
        return
    note = input_optional("Nhập ghi chú (có thể bỏ trống): ")

    contacts.append({"name": name, "phone": phone, "email": email, "note": note})
    save_contacts(contacts)
    print("✅ Đã thêm liên hệ!\n")


def edit_contact(contacts):
    phone = normalize_phone(input("Nhập số điện thoại cần sửa: ").strip())
    for c in contacts:
        if c['phone'] == phone:
            new_name = input_optional(f"Tên mới (hiện tại: {c['name']}): ", c['name'])
            new_phone = normalize_phone(input_optional(f"SĐT mới (hiện tại: {c['phone']}): ", c['phone']))
            if not validate_phone(new_phone):
                print("❌ Số điện thoại không hợp lệ!\n")
                return
            new_email = input_optional(f"Email mới (hiện tại: {c.get('email','')}): ", c.get("email", ""))
            if not validate_email(new_email):
                print("❌ Email không hợp lệ!\n")
                return
            new_note = input_optional(f"Ghi chú mới (hiện tại: {c.get('note','')}): ", c.get("note", ""))

            c.update({"name": new_name, "phone": new_phone, "email": new_email, "note": new_note})
            save_contacts(contacts)
            print("✅ Đã cập nhật liên hệ!\n")
            return
    print("❌ Không tìm thấy số này!\n")


def delete_contact(contacts):
    phone = normalize_phone(input("Nhập số điện thoại cần xóa: ").strip())
    for c in contacts:
        if c['phone'] == phone:
            contacts.remove(c)
            save_contacts(contacts)
            print("✅ Đã xóa liên hệ!\n")
            return
    print("❌ Không tìm thấy số này!\n")


def search_contact(contacts, export=False):
    keyword = input("Nhập tên/số/email cần tìm: ").strip().lower()
    results = [
        c for c in contacts
        if keyword in c['name'].lower()
        or keyword in c['phone']
        or keyword in c.get("email", "").lower()
        or keyword in c.get("note", "").lower()
    ]
    if results:
        print("\n📌 Kết quả tìm kiếm:")
        for c in results:
            print(f"- {c['name']} | {c['phone']} | {c.get('email','')} | {c.get('note','')}")
        print()

        if export:
            export_results(results)
    else:
        print("❌ Không tìm thấy!\n")


def list_contacts(contacts):
    if not contacts:
        print("❌ Danh bạ trống!\n")
        return

    sort_by = input("Sắp xếp theo (name/phone/email): ").strip().lower()
    if sort_by == "phone":
        contacts = sorted(contacts, key=lambda x: x['phone'])
    elif sort_by == "email":
        contacts = sorted(contacts, key=lambda x: x.get('email', "").lower())
    else:
        contacts = sorted(contacts, key=lambda x: x['name'].lower())

    total_pages = (len(contacts) + PAGE_SIZE - 1) // PAGE_SIZE
    page = 1

    while True:
        start = (page - 1) * PAGE_SIZE
        end = start + PAGE_SIZE
        print(f"\n📖 Trang {page}/{total_pages}")
        for c in contacts[start:end]:
            print(f"- {c['name']} | {c['phone']} | {c.get('email','')} | {c.get('note','')}")

        if page < total_pages:
            next_action = input("Nhấn Enter để sang trang tiếp (hoặc gõ q để thoát): ").strip().lower()
            if next_action == "q":
                break
            page += 1
        else:
            break
    print()


def import_contacts(contacts):
    filename = input("Nhập tên file JSON cần import: ").strip()
    if not os.path.exists(filename):
        print("❌ File không tồn tại!\n")
        return

    with open(filename, "r", encoding="utf-8") as f:
        new_contacts = json.load(f)

    for new_c in new_contacts:
        phone = normalize_phone(new_c['phone'])
        if not validate_phone(phone):
            print(f"❌ Bỏ qua số không hợp lệ: {new_c['phone']}")
            continue

        email = new_c.get("email", "")
        if not validate_email(email):
            print(f"❌ Bỏ qua email không hợp lệ: {email}")
            continue

        exists = next((c for c in contacts if c['phone'] == phone), None)
        if exists:
            choice = input(f"SĐT {phone} đã tồn tại ({exists['name']}). Cập nhật? (y/n): ").strip().lower()
            if choice == "y":
                exists.update({
                    "name": new_c.get("name", exists["name"]),
                    "phone": phone,
                    "email": email,
                    "note": new_c.get("note", exists.get("note",""))
                })
        else:
            contacts.append({
                "name": new_c.get("name", ""),
                "phone": phone,
                "email": email,
                "note": new_c.get("note", "")
            })

    save_contacts(contacts)
    print("✅ Import thành công!\n")


# ================== XUẤT FILE ==================

def export_results(results):
    print("Chọn định dạng xuất:")
    print("1) Word (.docx)")
    print("2) Excel (.xlsx)")
    choice = input("Chọn: ").strip()

    if choice == "1":
        filename = "ket_qua.docx"
        doc = Document()
        doc.add_heading("Kết quả tìm kiếm danh bạ", level=1)
        for c in results:
            doc.add_paragraph(f"{c['name']} | {c['phone']} | {c.get('email','')} | {c.get('note','')}")
        doc.save(filename)
        print(f"✅ Đã xuất ra file Word: {filename}\n")

    elif choice == "2":
        filename = "ket_qua.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(["Tên", "SĐT", "Email", "Ghi chú"])
        for c in results:
            ws.append([c['name'], c['phone'], c.get("email",""), c.get("note","")])
        wb.save(filename)
        print(f"✅ Đã xuất ra file Excel: {filename}\n")
    else:
        print("❌ Lựa chọn không hợp lệ!\n")


# ================== MENU ==================

def menu():
    print("=== QUẢN LÝ DANH BẠ ===")
    print("1) Thêm liên hệ")
    print("2) Sửa liên hệ")
    print("3) Xóa liên hệ")
    print("4) Tìm kiếm liên hệ")
    print("5) Hiển thị danh bạ (sắp xếp & phân trang)")
    print("6) Import từ file JSON")
    print("7) Tìm kiếm & xuất kết quả ra file (Word/Excel)")
    print("0) Thoát")


def main():
    contacts = load_contacts()
    while True:
        menu()
        choice = input("Chọn chức năng: ").strip()
        if choice == "1":
            add_contact(contacts)
        elif choice == "2":
            edit_contact(contacts)
        elif choice == "3":
            delete_contact(contacts)
        elif choice == "4":
            search_contact(contacts)
        elif choice == "5":
            list_contacts(contacts)
        elif choice == "6":
            import_contacts(contacts)
        elif choice == "7":
            search_contact(contacts, export=True)
        elif choice == "0":
            print("👋 Tạm biệt!")
            break
        else:
            print("❌ Lựa chọn không hợp lệ!\n")


if __name__ == "__main__":
    main()
